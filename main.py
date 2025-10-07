import os
import streamlit as st
import pandas as pd
from docx import Document
import io
import re
from typing import Tuple, List, Optional, Dict
import logging
from pathlib import Path
from llama_parse import LlamaParse
from dotenv import load_dotenv
import tempfile
from openai import AzureOpenAI
import base64
from PIL import Image
import fitz  # PyMuPDF for image extraction
import json

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Get API keys from environment variables
AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
LLAMA_CLOUD_API_KEY = os.getenv("LLAMA_CLOUD_API_KEY")

if not AZURE_OPENAI_KEY or not AZURE_OPENAI_ENDPOINT or not LLAMA_CLOUD_API_KEY:
    raise ValueError("Missing required API keys in environment variables")

# Initialize Azure OpenAI client
try:
    client = AzureOpenAI(
        api_key=AZURE_OPENAI_KEY,
        api_version=AZURE_OPENAI_API_VERSION,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        timeout=60.0,
        max_retries=3
    )
    logger.info(f"Azure OpenAI client initialized with endpoint: {AZURE_OPENAI_ENDPOINT}")
except Exception as e:
    logger.error(f"Failed to initialize Azure OpenAI client: {e}")
    raise

# Initialize LlamaParse with vision mode for better image handling
parser = LlamaParse(
    api_key=LLAMA_CLOUD_API_KEY,
    result_type="markdown",
    verbose=True,
    parse_images=True  # Enable image parsing
)

def extract_images_from_pdf(file) -> List[Dict]:
    """
    Extract images from PDF file using PyMuPDF
    
    Returns:
        List of dicts containing image data and metadata
    """
    images = []
    try:
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_path = tmp_file.name

        # Open PDF
        pdf_document = fitz.open(tmp_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Convert to PIL Image
                image = Image.open(io.BytesIO(image_bytes))
                
                # Only process reasonably sized images (likely diagrams, not logos)
                if image.width > 200 and image.height > 200:
                    # Convert to base64 for API
                    buffered = io.BytesIO()
                    image.save(buffered, format="PNG")
                    img_base64 = base64.b64encode(buffered.getvalue()).decode()
                    
                    images.append({
                        'page': page_num + 1,
                        'index': img_index,
                        'image': image,
                        'base64': img_base64,
                        'width': image.width,
                        'height': image.height
                    })
        
        pdf_document.close()
        os.unlink(tmp_path)
        
        logger.info(f"Extracted {len(images)} images from PDF")
        return images
        
    except Exception as e:
        logger.error(f"Image extraction failed: {e}")
        return []

def analyze_image_with_vision(image_base64: str, context: str = "") -> str:
    """
    Analyze image using Azure OpenAI Vision API (GPT-4 Vision)
    
    Args:
        image_base64: Base64 encoded image
        context: Optional context from surrounding text
        
    Returns:
        Description of the image
    """
    try:
        messages = [
            {
                "role": "system",
                "content": "You are an expert at analyzing technical diagrams, flowcharts, and architecture diagrams. Provide detailed descriptions including components, relationships, and flow."
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"""Analyze this diagram/image from an RFP document. 
                        
Context from document: {context if context else 'No additional context'}

Please provide:
1. Type of diagram (flowchart, architecture, network, etc.)
2. Main components and their relationships
3. Key flows or connections
4. Technical insights
5. Any text or labels visible in the diagram"""
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{image_base64}"
                        }
                    }
                ]
            }
        ]
        
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,  # Must be GPT-4 Vision model
            messages=messages,
            max_tokens=1000,
            temperature=0.3
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Vision API error: {e}")
        return f"Could not analyze image: {str(e)}"

def generate_mermaid_diagram(text: str, diagram_type: str = "flowchart") -> str:
    """
    Generate Mermaid diagram code from text description
    
    Args:
        text: Text describing the process/architecture
        diagram_type: Type of diagram to generate
        
    Returns:
        Mermaid diagram code
    """
    try:
        prompt = f"""Based on the following RFP content, generate a Mermaid {diagram_type} diagram.

Content:
{text}

Generate ONLY the Mermaid code. Use appropriate syntax:
- For flowchart: flowchart TD
- For sequence: sequenceDiagram
- For architecture: graph LR
- Strictly follow the mermaid.js syntax and do not use curly braces in middle

Make it clear and readable with proper labels."""

        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are an expert at creating Mermaid diagrams. Generate only valid Mermaid code."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.5
        )
        
        # Extract mermaid code from response
        mermaid_code = response.choices[0].message.content
        
        # Clean up the response to get just the mermaid code
        if "```mermaid" in mermaid_code:
            mermaid_code = mermaid_code.split("```mermaid")[1].split("```")[0].strip()
        elif "```" in mermaid_code:
            mermaid_code = mermaid_code.split("```")[1].split("```")[0].strip()
            
        return mermaid_code
        
    except Exception as e:
        logger.error(f"Mermaid generation error: {e}")
        return ""

def extract_text_from_pdf(file) -> Tuple[str, List[pd.DataFrame], List[Dict]]:
    """
    Extract text, tables, and images from PDF file
    
    Returns:
        Tuple[str, List[pd.DataFrame], List[Dict]]: Extracted text, tables, and images
    """
    try:
        # Extract text and tables using LlamaParse
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_path = tmp_file.name

        result = parser.load_data(tmp_path)
        text = "\n\n---\n\n".join([page.text for page in result])
        
        # Extract tables
        tables = []
        for page in result:
            if hasattr(page, 'tables') and page.tables:
                for table in page.tables:
                    try:
                        df = pd.DataFrame(table)
                        if not df.empty:
                            tables.append(df)
                    except Exception as e:
                        logger.warning(f"Failed to convert table to DataFrame: {e}")

        os.unlink(tmp_path)
        
        # Extract images separately
        file.seek(0)  # Reset file pointer
        images = extract_images_from_pdf(file)
            
        return text, tables, images
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Failed to process PDF: {str(e)}")

def extract_text_from_docx(file):
    """Extract text, tables, and images from DOCX file"""
    doc = Document(file)
    text = ""
    tables = []
    images = []
    
    # Extract text
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        if len(table_data) > 1:
            tables.append(pd.DataFrame(table_data[1:], columns=table_data[0]))
    
    # Extract images from DOCX
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data))
                
                if image.width > 200 and image.height > 200:
                    buffered = io.BytesIO()
                    image.save(buffered, format="PNG")
                    img_base64 = base64.b64encode(buffered.getvalue()).decode()
                    
                    images.append({
                        'page': 0,
                        'index': len(images),
                        'image': image,
                        'base64': img_base64,
                        'width': image.width,
                        'height': image.height
                    })
            except Exception as e:
                logger.warning(f"Failed to extract image from DOCX: {e}")
    
    return text, tables, images

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    text = file.getvalue().decode()
    return text, [], []

def clean_text(text):
    """Clean extracted text"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    return text.strip()

def call_azure_openai(prompt: str, max_tokens: int = 4000) -> str:
    """Call Azure OpenAI API with the given prompt"""
    try:
        logger.info(f"Calling Azure OpenAI with deployment: {AZURE_OPENAI_DEPLOYMENT}")
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are an expert RFP analyst. Analyze documents thoroughly and provide detailed, structured insights."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens,
            temperature=0.7,
            top_p=0.95,
            timeout=120
        )
        logger.info("Azure OpenAI response received successfully")
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Azure OpenAI API error: {e}")
        raise RuntimeError(f"Failed to get AI response: {str(e)}")

def analyze_section(text, section_type="requirements"):
    """Analyze specific sections of the RFP"""
    prompts = {
        "requirements": """
            Analyze the technical requirements in this section:
            1. List all mandatory requirements
            2. Identify optional requirements
            3. Flag any unclear specifications
            4. Note any technical constraints
            """,
        "timeline": """
            Analyze the timeline aspects:
            1. List all key dates and deadlines
            2. Identify major milestones
            3. Flag any tight or unrealistic timelines
            4. Note dependencies between phases
            """,
        "compliance": """
            Analyze compliance requirements:
            1. List all regulatory requirements
            2. Identify certification needs
            3. Note security requirements
            4. Flag critical compliance issues
            """
    }
    
    full_prompt = f"{prompts.get(section_type, prompts['requirements'])}\n\nContent:\n{text}"
    return call_azure_openai(full_prompt)

def analyze_tables(tables):
    """Analyze tables found in the document"""
    analysis = []
    for i, table in enumerate(tables):
        if not table.empty:
            analysis.append(f"\nTable {i+1} Analysis:")
            analysis.append(f"- Columns: {', '.join(table.columns.tolist())}")
            analysis.append(f"- Rows: {len(table)}")
            analysis.append("- Content Summary:")
            for col in table.columns:
                unique_values = table[col].nunique()
                analysis.append(f"  * {col}: {unique_values} unique values")
    
    return "\n".join(analysis)

def analyze_rfp_content(text: str) -> str:
    """Analyze the full content of the RFP document."""
    prompt = f"""
    Analyze this RFP document and provide a detailed markdown-formatted response covering:

    # Executive Summary
    - Brief overview of the RFP
    - Key objectives and goals

    # Requirements Analysis
    - ## Technical Requirements
      - Core technical specifications
      - Optional features
      - Integration requirements
    - ## Business Requirements
      - Mandatory business needs
      - Optional enhancements
    
    # Timeline and Milestones
    - Key dates and deadlines
    - Project phases
    - Dependencies

    # Risk Assessment
    - Technical risks
    - Business risks
    - Compliance concerns

    # Recommendations
    - Strategic approach
    - Key focus areas
    - Potential challenges

    Please analyze the following content:
    {text}
    """
    try:
        return call_azure_openai(prompt, max_tokens=4000)
    except Exception as e:
        logger.error(f"Error analyzing RFP content: {e}")
        return "Error analyzing content. Please try again."

@st.cache_data
def process_document(file_content, file_type: str) -> Tuple[str, List[pd.DataFrame], List[Dict]]:
    """Cache the document processing results"""
    if file_type == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_type == 'docx':
        return extract_text_from_docx(file_content)
    else:  # txt
        return extract_text_from_txt(file_content)

def main():
    st.set_page_config(page_title="RFP Document Analyzer", layout="wide")
    
    if not AZURE_OPENAI_KEY or not AZURE_OPENAI_ENDPOINT or not LLAMA_CLOUD_API_KEY:
        st.error("Missing required API keys in environment variables!")
        st.stop()
    
    with st.sidebar:
        st.title("📄 RFP Document Analyzer")
        st.markdown("Upload your RFP document for analysis.")
        st.markdown("*Powered by Azure OpenAI + Vision*")
        
        MAX_FILE_SIZE = 50 * 1024 * 1024
        uploaded_file = st.file_uploader(
            "Choose an RFP document",
            type=['pdf', 'docx', 'txt'],
            help="Upload a PDF, DOCX, or TXT file (max 50MB)"
        )
        
        if uploaded_file:
            st.write(f"File: {uploaded_file.name}")
            st.write(f"Size: {uploaded_file.size/1024:.2f} KB")
            
            st.markdown("### 📊 Analysis Options")
            analyze_full = st.button("🔍 Full Analysis", use_container_width=True)
            analyze_req = st.button("📋 Requirements", use_container_width=True)
            analyze_timeline = st.button("⏱️ Timeline", use_container_width=True)
            analyze_compliance = st.button("✓ Compliance", use_container_width=True)
            
            st.markdown("### 🎨 Diagram Generation")
            generate_flow = st.button("📊 Generate Flowchart", use_container_width=True)
            generate_arch = st.button("🏗️ Generate Architecture", use_container_width=True)
            analyze_images = st.button("🖼️ Analyze Images", use_container_width=True)
        
        st.markdown("---")
        st.markdown("### 💡 Features")
        st.markdown("""
        - Text & table extraction
        - Image & diagram analysis
        - Auto diagram generation
        - Vision AI analysis
        """)

    if uploaded_file:
        if uploaded_file.size > MAX_FILE_SIZE:
            st.error("File size exceeds 50MB limit.")
            return
            
        try:
            text, tables, images = process_document(
                uploaded_file,
                uploaded_file.name.split('.')[-1].lower()
            )
            
            cleaned_text = clean_text(text)
            
            # Create tabs for better organization
            tab1, tab2, tab3, tab4 = st.tabs(["📑 Analysis", "🖼️ Images", "📊 Diagrams", "📄 Raw Data"])
            
            with tab1:
                st.markdown("### 📑 Analysis Results")
                
                if analyze_full:
                    with st.spinner("Analyzing full content..."):
                        analysis = analyze_rfp_content(cleaned_text)
                        st.markdown(analysis)
                
                if analyze_req:
                    with st.spinner("Analyzing requirements..."):
                        req_analysis = analyze_section(cleaned_text, "requirements")
                        st.markdown(req_analysis)
                
                if analyze_timeline:
                    with st.spinner("Analyzing timeline..."):
                        timeline_analysis = analyze_section(cleaned_text, "timeline")
                        st.markdown(timeline_analysis)
                
                if analyze_compliance:
                    with st.spinner("Analyzing compliance..."):
                        compliance_analysis = analyze_section(cleaned_text, "compliance")
                        st.markdown(compliance_analysis)
            
            with tab2:
                st.markdown("### 🖼️ Extracted Images & Diagrams")
                
                if images:
                    st.info(f"Found {len(images)} images in the document")
                    
                    if analyze_images:
                        for idx, img_data in enumerate(images):
                            with st.expander(f"Image {idx+1} (Page {img_data['page']})", expanded=True):
                                col1, col2 = st.columns([1, 2])
                                
                                with col1:
                                    st.image(img_data['image'], caption=f"Image {idx+1}", use_container_width=True)
                                    st.caption(f"Size: {img_data['width']}x{img_data['height']}px")
                                
                                with col2:
                                    with st.spinner("Analyzing image with AI..."):
                                        # Get surrounding context (simplified)
                                        context = cleaned_text[:500]
                                        description = analyze_image_with_vision(img_data['base64'], context)
                                        st.markdown("**AI Analysis:**")
                                        st.markdown(description)
                    else:
                        for idx, img_data in enumerate(images):
                            with st.expander(f"Image {idx+1} (Page {img_data['page']})"):
                                st.image(img_data['image'], use_container_width=True)
                                st.caption(f"Size: {img_data['width']}x{img_data['height']}px")
                                st.info("Click 'Analyze Images' button to get AI analysis")
                else:
                    st.info("No images found in the document")
            
            with tab3:
                st.markdown("### 📊 Generated Diagrams")
                
                if generate_flow:
                    with st.spinner("Generating flowchart..."):
                        mermaid_code = generate_mermaid_diagram(cleaned_text[:2000], "flowchart")
                        if mermaid_code:
                            st.markdown("**Generated Flowchart:**")
                            st.code(mermaid_code, language="mermaid")
                            # Note: Streamlit doesn't natively render Mermaid, but you can use st.components
                            st.info("Copy the code above and paste it into https://mermaid.live to visualize")
                
                if generate_arch:
                    with st.spinner("Generating architecture diagram..."):
                        mermaid_code = generate_mermaid_diagram(cleaned_text[:2000], "architecture")
                        if mermaid_code:
                            st.markdown("**Generated Architecture Diagram:**")
                            st.code(mermaid_code, language="mermaid")
                            st.info("Copy the code above and paste it into https://mermaid.live to visualize")
            
            with tab4:
                col1, col2 = st.columns(2)
                
                with col1:
                    with st.expander("📄 Extracted Text", expanded=False):
                        st.text_area("Content", cleaned_text, height=400)
                
                with col2:
                    if tables:
                        with st.expander("📊 Extracted Tables", expanded=False):
                            table_analysis = analyze_tables(tables)
                            st.markdown(table_analysis)
                            for i, table in enumerate(tables):
                                st.markdown(f"**Table {i+1}**")
                                st.dataframe(table, use_container_width=True)
                
        except Exception as e:
            st.error(f"Error processing document: {str(e)}")
            logger.exception("Document processing error")
    else:
        st.markdown("""
        # Welcome to RFP Document Analyzer! 👋
        
        ## Features:
        
        ### 📄 Document Analysis
        - Extract and analyze text content
        - Identify requirements, timelines, compliance
        
        ### 🖼️ Image Processing
        - Extract diagrams and flowcharts from PDFs
        - AI-powered image analysis with GPT-4 Vision
        - Understand architecture diagrams automatically
        
        ### 📊 Diagram Generation
        - Auto-generate flowcharts from text
        - Create architecture diagrams
        - Export as Mermaid code
        
        ### 🚀 Get Started
        Upload your RFP document using the sidebar to begin!
        """)

if __name__ == "__main__":
    main()